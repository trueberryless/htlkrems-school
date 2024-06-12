#!/usr/bin/env python
# coding: utf-8

# In[1]:


import PyPDF2 as pypdf


# In[2]:


with open('Cyber-Security.pdf', 'rb') as pdf:
    reader = pypdf.PdfReader(pdf)
    metadata = reader.metadata
    
    print("Author: " + metadata.author)
    print("Producter: " + metadata.producer)
    print("Title: " + metadata.title)


# In[3]:


# 2.2

with open('Cyber-Security.pdf', 'rb') as pdf:
    reader = pypdf.PdfReader(pdf)
    print("erste Seite:", reader.pages[0].extract_text())


# In[4]:


with open('Cyber-Security.pdf', 'rb') as pdf:
    reader = pypdf.PdfReader(pdf)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    print(text)


# In[5]:


with open('Cyber-Security.pdf', 'rb') as pdf:
    reader = pypdf.PdfReader(pdf)
    pages = []
    for page in reader.pages:
        pages.append(page)
    print(len(pages))


# In[8]:


get_ipython().system('pip install python-docx')

from docx import Document
from docx.shared import Inches
document = Document()
document.add_heading('Document Title', 0)
p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')
document.add_paragraph(
 'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
 'first item in ordered list', style='List Number'
)
document.add_picture('trueberryless_under1MB.jpg', width=Inches(1.25))
records = (
 (3, '101', 'Spam'),
 (7, '422', 'Eggs'),
 (4, '631', 'Spam, spam, eggs, and spam')
)
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
 row_cells = table.add_row().cells
 row_cells[0].text = str(qty)
 row_cells[1].text = id
 row_cells[2].text = desc
document.add_page_break()
document.save('demo.docx') 


# ![image.png](attachment:image.png)

# In[9]:


# Bibliotheken laden
import urllib.request as urllib2
from bs4 import BeautifulSoup
# Website laden
response = urllib2.urlopen('https://www.htlkrems.ac.at')
html_doc = response.read()
# Das BeautifulSoup Object soup repräsentiert das „geparste“ HTML-Dokument
soup = BeautifulSoup(html_doc, 'html.parser')
# Das „geparste“ HTML-Dokument formatieren, sodass jeder Tag bzw. Textblock
# in einer separaten Zeile ausgegeben wird
strhtm = soup.prettify()
# Ein paar Zeilen ausgeben
print (strhtm[:1000]) 


# In[10]:


# 2.3
a_list = soup.find_all('a')

for a in a_list:
    print(a.get_text())


# In[11]:


# 2.4
len(a_list)

